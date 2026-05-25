from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = "personal_data_consent_merch_bot.docx"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(document, text, level=1):
    paragraph = document.add_heading(level=level)
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True, color=(31, 78, 121))
    return paragraph


def add_body_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_number(document, text):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def build_document():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.text = "Мерч Бот | обработка персональных данных"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color=(100, 100, 100))

    footer = section.footer.paragraphs[0]
    footer.text = "Документ для ознакомления пользователя перед передачей данных"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color=(100, 100, 100))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Согласие на обработку персональных данных")
    set_run_font(title_run, size=20, bold=True, color=(31, 78, 121))

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("для пользователей Мерч Бота")
    set_run_font(subtitle_run, size=12, color=(90, 90, 90))

    add_body_paragraph(
        document,
        "Настоящий документ предназначен для направления пользователю перед сбором "
        "контактных данных в Мерч Боте. Перед использованием рекомендуется дополнить "
        "документ реквизитами оператора персональных данных и ссылкой на действующую "
        "политику обработки персональных данных.",
    )

    add_heading(document, "1. Данные, которые запрашивает бот", 1)
    add_body_paragraph(
        document,
        "Для подбора вакансий и связи с пользователем бот может запросить и сохранить:",
    )
    add_bullet(document, "город пользователя;")
    add_bullet(document, "фамилию, имя и отчество;")
    add_bullet(document, "номер телефона для связи;")
    add_bullet(document, "информацию об отклике на выбранную вакансию.")

    add_heading(document, "2. Цели обработки", 1)
    add_bullet(document, "идентификация пользователя в переписке с ботом;")
    add_bullet(document, "подбор подходящих вакансий;")
    add_bullet(document, "передача отклика специалисту по подбору;")
    add_bullet(document, "обратная связь по вакансии и уточнение деталей трудоустройства;")
    add_bullet(document, "ведение внутреннего учета откликов.")

    add_heading(document, "3. Действия с персональными данными", 1)
    add_body_paragraph(
        document,
        "С персональными данными могут выполняться следующие действия: сбор, запись, "
        "систематизация, накопление, хранение, уточнение, использование, передача "
        "уполномоченным сотрудникам, удаление и уничтожение.",
    )

    add_heading(document, "4. Хранение и удаление данных", 1)
    add_body_paragraph(
        document,
        "Данные хранятся в информационных системах оператора и используются только для "
        "целей, связанных с подбором вакансий и обработкой откликов. Пользователь может "
        "посмотреть, скорректировать или удалить свои данные через кнопку «Мои данные» в боте.",
    )

    add_heading(document, "5. Права пользователя", 1)
    add_body_paragraph(
        document,
        "Пользователь вправе запросить сведения об обработке персональных данных, "
        "потребовать их уточнения, блокирования или удаления, а также отозвать согласие "
        "на обработку персональных данных.",
    )

    add_heading(document, "6. Текст согласия", 1)
    add_body_paragraph(
        document,
        "Нажимая кнопку «Согласен» в боте, пользователь подтверждает:",
    )
    add_number(
        document,
        "Я даю согласие на обработку моих персональных данных: города, ФИО и телефона "
        "для подбора вакансий, связи со мной и фиксации откликов.",
    )
    add_number(
        document,
        "Я уведомлен(а), что могу посмотреть, изменить или удалить свои данные через "
        "раздел «Мои данные» в боте.",
    )
    add_number(
        document,
        "Я понимаю, что отзыв согласия может ограничить возможность отправлять отклики "
        "на вакансии через бот.",
    )

    add_heading(document, "7. Реквизиты оператора", 1)
    table = document.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Оператор", "Указать наименование компании / ИП"),
        ("Контакт для обращений", "Указать email или телефон"),
        ("Политика обработки ПДн", "Указать ссылку на политику"),
        ("Дата редакции", "25.05.2026"),
    ]
    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=10)

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
